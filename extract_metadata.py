# -*- coding: utf-8 -*-
import os
import re
import json

try:
    import fitz  # PyMuPDF
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# File mapping: user id -> filename pattern
FILE_MAP = {
    1: "1. CONSENTIMIENTO INFORMADO CIRUGIA BARIATRICA.docx",
    2: "2. CONSENTIMIENTO INFORMADO DE CIRUGÍA CARDIOVASCULAR.pdf",
    3: "3. CONSENTIMIENTO INFORMADO PARA ACOMPAÑANTE DE CASOS PROBABLE CONFIRMADO DE COVID-19.pdf",
    4: "4. CONSENTIMIENTO INFORMADO PARA LA ADMINISTRACION DE COMPONENTES SANGUINEOS.pdf",
    5: "5. CONSENTIMIENTO INFORMADO PARA LA REALIZACION DE ESTUDIOS DE INVESTIGACIÓN.pdf",
    6: "6. CONSENTIMIENTO INFORMADO PARA REALIZACIÓN DE PROCEDIMIENTOS DURANTE LA PANDEMIA DEL NUEVO CORONAVIRUS (COVID-19).pdf",
    7: "7. CONSENTIMIENTO INFORMADO PARA REALIZAR LA PRUEBA DE VIH.pdf",
    8: "8. CONSENTIMIENTO INFORMADO PARA VISITA DOMICILIARIA.docx",
    9: "9. CONSENTIMIENTO INFORMADO PSICOLOGIA.docx",
    10: "10. CONSENTIMIENTO INFORMADO DE ANESTESIA.pdf",
    12: "12. CONSENTIMIENTO INFORMADO MÉDICO.docx",
    13: "13. CONSENTIMIENTO INFORMADO DE FISIOTERAPIA.docx",
    14: "14. CONSENTIMIENTO INFORMADO PARA LA OBTENCION DE IMAGENES CON CONTENIDO CLINICO.doc",
    16: "16. CONSENTIMIENTO INFORMADO DE ENFERMERÍA.pdf",
}

DOCS_DIR = r"c:\Users\sistemas.practicante\Downloads\CONSENTIMIENTOS INFORMADOS\CONSENTIMIENTOS INFORMADOS\documentos"

def parse_metadata(text):
    codigo = version = vigente = None
    if not text:
        return codigo, version, vigente
    # Código: F-GQ-XXX, F-CX-CCM-XXX, etc.
    codigo_match = re.search(r'C[oó]DIGO\s*:?\s*([Ff]-[A-Za-z-]+-\d+)', text, re.IGNORECASE)
    if codigo_match:
        codigo = codigo_match.group(1).upper().replace(' ', '')
    if not codigo:
        codigo_match = re.search(r'[Ff]\s*-\s*GQ\s*-\s*(\d+)|F-GQ-(\d+)|F\s*GQ\s*(\d+)', text, re.IGNORECASE)
        if codigo_match:
            num = codigo_match.group(1) or codigo_match.group(2) or codigo_match.group(3)
            codigo = f"F-GQ-{num}"
    if not codigo:
        codigo_match = re.search(r'[Ff]-[A-Za-z]+-[A-Za-z]*-(\d+)|[Ff]-[A-Za-z]+-(\d+)', text, re.IGNORECASE)
        if codigo_match:
            parts = re.search(r'([Ff]-[A-Za-z-]+-\d+)', text, re.IGNORECASE)
            if parts:
                codigo = parts.group(1).upper()
    # Versión: N.° X, VERSION: X, V1, etc. (reject 0 as likely false match)
    ver_match = re.search(r'Versi[oó]n\s*N\.?\s*[°º]?\s*(\d+)', text, re.IGNORECASE)
    if ver_match and ver_match.group(1) != '0':
        version = f"N.\u00b0 {ver_match.group(1)}"
    if not version:
        ver_match = re.search(r'VERSION\s*:?\s*(\d+)', text, re.IGNORECASE)
        if ver_match:
            version = f"V{ver_match.group(1)}"
    if not version:
        ver_match = re.search(r'\bV\s*(\d+)\b', text)
        if ver_match and ver_match.group(1) != '0':
            version = f"V{ver_match.group(1)}"
    # Vigente: Month Year or VIGENTE DESDE: DD-MM-YYYY
    vig_match = re.search(r'Vigente\s*:?\s*([A-Za-záéíóúñ]+)\s+(\d{4})', text, re.IGNORECASE)
    if vig_match:
        vigente = f"{vig_match.group(1)} {vig_match.group(2)}"
    if not vigente:
        text_norm = re.sub(r'\s+', ' ', text)
        vig_match = re.search(r'VIGENTE DESDE\s*:?\s*(\d{1,2})-(\d{1,2})-\s*(\d{4})', text_norm, re.IGNORECASE)
        if vig_match:
            months = ['Enero','Febrero','Marzo','Abril','Mayo','Junio','Julio','Agosto','Septiembre','Octubre','Noviembre','Diciembre']
            d, m, y = int(vig_match.group(1)), int(vig_match.group(2)), vig_match.group(3)
            if 1 <= d <= 12 and 1 <= m <= 12:
                vigente = f"{months[m-1]} {y}"  # assume DD-MM-YYYY
            else:
                vigente = f"{months[d-1]} {y}" if 1 <= d <= 12 else f"{months[m-1]} {y}"
    if not vigente:
        vig_match = re.search(r'Vigente\s*:?\s*([A-Za-záéíóúñ]+\s+\d{4})', text, re.IGNORECASE)
        if vig_match:
            vigente = vig_match.group(1).strip()
    return codigo, version, vigente

def extract_from_pdf(filepath):
    codigo = version = vigente = None
    texts = []
    # Try PyMuPDF first (often better for complex PDFs)
    if HAS_FITZ:
        try:
            doc = fitz.open(filepath)
            for i in range(min(5, len(doc))):
                texts.append(doc[i].get_text())
            doc.close()
        except Exception:
            pass
    # Fallback to pdfplumber, also try table extraction
    if not texts and HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages[:5]:
                    t = page.extract_text()
                    if t:
                        texts.append(t)
                    for table in page.extract_tables() or []:
                        for row in table:
                            texts.append(" ".join(str(c) for c in row if c))
        except Exception:
            pass
    text = "\n".join(texts)
    codigo, version, vigente = parse_metadata(text)
    return codigo, version, vigente

def extract_from_docx(filepath):
    if not HAS_DOCX:
        return None, None, None
    try:
        doc = DocxDocument(filepath)
        text = "\n".join(p.text for p in doc.paragraphs) + "\n"
        for table in doc.tables:
            for row in table.rows:
                text += " ".join(cell.text for cell in row.cells) + "\n"
        return parse_metadata(text)
    except Exception:
        return None, None, None

def main():
    results = {}
    all_files = os.listdir(DOCS_DIR)
    for doc_id, filename in FILE_MAP.items():
        filepath = None
        prefix = str(doc_id) + "."
        for f in all_files:
            if f.startswith(prefix) and f.lower().endswith(('.pdf', '.docx', '.doc')):
                filepath = os.path.join(DOCS_DIR, f)
                break
        if not filepath or not os.path.exists(filepath):
            results[doc_id] = {"codigo": "unknown", "version": "unknown", "vigente": "unknown"}
            continue
            
        if filepath.lower().endswith('.pdf'):
            codigo, version, vigente = extract_from_pdf(filepath)
        elif filepath.lower().endswith('.docx'):
            codigo, version, vigente = extract_from_docx(filepath)
        else:
            codigo, version, vigente = None, None, None
        # Override: anesthesia (id 10) is image-based PDF - user provided from image
        if doc_id == 10 and not codigo:
            codigo, version, vigente = "F-GQ-129", "V1", "Abril 2020"
        results[doc_id] = {
            "codigo": codigo or "unknown",
            "version": version or "unknown",
            "vigente": vigente or "unknown"
        }
    
    print(json.dumps(results, indent=2, ensure_ascii=False))

if __name__ == "__main__":
    main()
